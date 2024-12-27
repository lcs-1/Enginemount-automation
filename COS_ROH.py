import tkinter as tk
from tkinter import messagebox, filedialog
import pdfplumber
from datetime import datetime
from docxtpl import DocxTemplate
from docx.shared import Inches
import os

def save_document(event=None):
    co_name = co_name_entry.get()
    part = part_entry.get()
    p_no = p_no_entry.get()
    p_sn = p_sn_entry.get()
    date = date_entry.get()
    ac_type = ac_type_entry.get()
    
    # Extract the first five digits of the Part Number (P No)
    first_five_digits = p_no[1:5]
    
    # Determine the engine based on the first five digits of P No and AC/type
    if first_five_digits == '7121':
        if ac_type == 'A330-800':
            engine = 'RR'
        else:
            engine = 'PW'
    elif first_five_digits == '7122':
        if ac_type != 'A330-800':
            engine = 'CFM'
        else:
            engine = ''  # No engine specified for this condition
    else:
        engine = ''  # No engine specified for other conditions
    
    # Determine the structure based on the first five digits of P No
    if first_five_digits == '7121':
        structure = '71-21'
    elif first_five_digits == '7122':
        structure = '71-22'
    else:
        structure = ''

    if not co_name or not part or not p_no or not p_sn or not date or not ac_type or not engine or not structure:
        messagebox.showerror("Error", "All fields must be filled.")
        return

    context = {'co_name': co_name, 'part': part, 'p_no': p_no, 'p_sn': p_sn, 'engine': engine, 'date': date, 'ac_type': ac_type, 'structure': structure}

    doc = DocxTemplate(r"C:\L.C.S\auto\Template-COS_damage.docx")
    doc.render(context)
    
    output_filename = f"COS-{co_name}_Design.docx"
    output_path = fr"C:\L.C.S\auto\{output_filename}"
    
    # Inserting PDF pages as images into the document starting from page 2
    pdf_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if pdf_path:
        print("PDF Path:", pdf_path)  # Debug print
        
        # Convert PDF pages to images, excluding the first, second, and last pages
        with pdfplumber.open(pdf_path) as pdf:
            images = []
            num_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                if i == 0 or i == 1 or i == num_pages - 1:
                    continue  # Skip the first, second, and last pages
                img_path = f"page_{i}.png"
                cropped_page = page.within_bbox((0, 100, page.width, page.height))
                cropped_page.to_image(resolution=300).save(img_path)  # Adjust crop parameters as needed
                print("Image saved:", img_path)  # Debug print
                images.append(img_path)
        
        # Add images to the document
        for i, img_path in enumerate(images):
            if i > 0:
                # Add a section break before inserting images except for the first page
                doc.add_section()
                
            if i == 0:
                # Insert the first image directly into the document and center it horizontally
                paragraph = doc.add_paragraph()
                paragraph.alignment = 1  # Set paragraph alignment to center
                run = paragraph.add_run()
                # Adjust the width of the image here
                run.add_picture(img_path, width=Inches(8))  # Adjust width as needed
                print("Image inserted into document:", img_path)  # Debug print
                os.remove(img_path)  # Remove the temporary image file
            else:
                # Insert other images in new pages
                paragraph = doc.add_paragraph()
                paragraph.alignment = 1  # Set paragraph alignment to center
                run = paragraph.add_run()
                # Adjust the width of the image here
                run.add_picture(img_path, width=Inches(8))  # Adjust width as needed
                print("Image inserted into document:", img_path)  # Debug print
                os.remove(img_path)  # Remove the temporary image file
    
    doc.save(output_path)
    messagebox.showinfo("Success", "Document saved successfully.")

def open_pdf_viewer(event=None):
    pdf_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if pdf_path:
        with pdfplumber.open(pdf_path) as pdf:
            first_page = pdf.pages[0]
            extracted_text = first_page.extract_text()
            text_dialog = tk.Toplevel(root)
            text_dialog.title("Selected Text from PDF")
            text_widget = tk.Text(text_dialog)
            text_widget.insert(tk.END, extracted_text)
            text_widget.pack()
            fill_form_from_text(extracted_text)

def on_hover_enter(event):
    if event.widget == pdf_viewer_button:
        event.widget.config(bg='#4D92FF', cursor='hand2')
    elif event.widget == save_button:
        event.widget.config(bg='#98fb98', cursor='hand2')

def on_hover_leave(event):
    if event.widget == pdf_viewer_button:
        event.widget.config(bg='white', cursor='')
    elif event.widget == save_button:
        event.widget.config(bg='white', cursor='')

def on_resize(event):
    label_width = max(len("CO Name:"), len("Part:"), len("P No:"), len("P SN:"), len("Engine:"), len("Date:"), len("AC Type:"), len("Structure:"))
    entry_width = 25
    padding_x = 10
    padding_y = 5

    co_name_label.config(width=label_width, font=('Arial', 12))
    part_label.config(width=label_width, font=('Arial', 12))
    p_no_label.config(width=label_width, font=('Arial', 12))
    p_sn_label.config(width=label_width, font=('Arial', 12))
    engine_label.config(width=label_width, font=('Arial', 12))
    date_label.config(width=label_width, font=('Arial', 12))
    ac_type_label.config(width=label_width, font=('Arial', 12))
    structure_label.config(width=label_width, font=('Arial', 12))

    co_name_entry.config(width=entry_width, font=('Arial', 12))
    part_entry.config(width=entry_width, font=('Arial', 12))
    p_no_entry.config(width=entry_width, font=('Arial', 12))
    p_sn_entry.config(width=entry_width, font=('Arial', 12))
    engine_entry.config(width=entry_width, font=('Arial', 12))
    date_entry.config(width=entry_width, font=('Arial', 12))
    ac_type_entry.config(width=entry_width, font=('Arial', 12))
    structure_entry.config(width=entry_width, font=('Arial', 12))

    pdf_viewer_button.config(width=label_width + entry_width + padding_x, font=('Arial', 12))
    save_button.config(width=label_width + entry_width + padding_x, font=('Arial', 12))

def fill_form_from_text(extracted_text):
    lines = extracted_text.split('\n')
    fields = {'Reference': co_name_entry}
    part_desc = ''
    part_sn = ''
    ac_type = ''
    for line in lines:
        for keyword, entry in fields.items():
            if keyword in line:
                entry.insert(tk.END, line.replace(keyword, '').strip())
                break
        
        if 'Part N°' in line:
            part_no = line.split('Part N°')[-1].strip().split()[0]
            p_no_entry.insert(tk.END, part_no)
        
        if 'Part Serial N°' in line:
            part_sn = line.split('Part Serial N°')[-1].strip().split('CA code')[0].strip()
            p_sn_entry.insert(tk.END, part_sn)
        
        if 'A/C Type' in line and not ac_type:
            ac_type = line.split('A/C Type')[-1].strip()
            ac_type_entry.insert(tk.END, ac_type)
            
        if 'Part Description' in line:
            part_desc = line.split('Part Description')[-1].strip()
            if 'A/C Type' in part_desc:
                part_desc, ac_type = map(str.strip, part_desc.split('A/C Type'))
                part_entry.insert(tk.END, part_desc)
                if not ac_type_entry.get():
                    ac_type_entry.insert(tk.END, ac_type)
            else:
                part_entry.insert(tk.END, part_desc)
    
    # Fill the date field with the current date in dd/mm/yyyy format
    today = datetime.now().strftime('%d/%m/%Y')
    date_entry.insert(tk.END, today)

root = tk.Tk()
root.title("Document Generator")

root.configure(bg='#F0F0F0')

co_name_label = tk.Label(root, text="CO Name:", bg='#F0F0F0')
co_name_entry = tk.Entry(root)

part_label = tk.Label(root, text="Part:", bg='#F0F0F0')
part_entry = tk.Entry(root)

p_no_label = tk.Label(root, text="P No:", bg='#F0F0F0')
p_no_entry = tk.Entry(root)

p_sn_label = tk.Label(root, text="P SN:", bg='#F0F0F0')
p_sn_entry = tk.Entry(root)

engine_label = tk.Label(root, text="Engine:", bg='#F0F0F0')
engine_entry = tk.Entry(root)

date_label = tk.Label(root, text="Date:", bg='#F0F0F0')
date_entry = tk.Entry(root)

ac_type_label = tk.Label(root, text="AC Type:", bg='#F0F0F0')
ac_type_entry = tk.Entry(root)

structure_label = tk.Label(root, text="Structure:", bg='#F0F0F0')
structure_entry = tk.Entry(root)

pdf_viewer_button = tk.Button(root, text="Open PDF Viewer", command=open_pdf_viewer, bg='white', activebackground='#4D92FF')
save_button = tk.Button(root, text="Save Document", command=save_document, bg='white', activebackground='#98fb98')

pdf_viewer_button.bind("<Enter>", on_hover_enter)
pdf_viewer_button.bind("<Leave>", on_hover_leave)
save_button.bind("<Enter>", on_hover_enter)
save_button.bind("<Leave>", on_hover_leave)

co_name_label.grid(row=0, column=0, padx=10, pady=5, sticky='e')
co_name_entry.grid(row=0, column=1, padx=10, pady=5)

part_label.grid(row=1, column=0, padx=10, pady=5, sticky='e')
part_entry.grid(row=1, column=1, padx=10, pady=5)

p_no_label.grid(row=2, column=0, padx=10, pady=5, sticky='e')
p_no_entry.grid(row=2, column=1, padx=10, pady=5)

p_sn_label.grid(row=3, column=0, padx=10, pady=5, sticky='e')
p_sn_entry.grid(row=3, column=1, padx=10, pady=5)

engine_label.grid(row=4, column=0, padx=10, pady=5, sticky='e')
engine_entry.grid(row=4, column=1, padx=10, pady=5)

date_label.grid(row=5, column=0, padx=10, pady=5, sticky='e')
date_entry.grid(row=5, column=1, padx=10, pady=5)

ac_type_label.grid(row=6, column=0, padx=10, pady=5, sticky='e')
ac_type_entry.grid(row=6, column=1, padx=10, pady=5)

structure_label.grid(row=7, column=0, padx=10, pady=5, sticky='e')
structure_entry.grid(row=7, column=1, padx=10, pady=5)

pdf_viewer_button.grid(row=8, column=0, columnspan=2, padx=10, pady=10)
save_button.grid(row=9, column=0, columnspan=2, padx=10, pady=10)

root.grid_columnconfigure(1, weight=1)
root.bind("<Configure>", on_resize)

root.mainloop()
